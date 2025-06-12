import re
import docx
import nltk
import numpy as np
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.decomposition import LatentDirichletAllocation
from nltk.corpus import stopwords

nltk.download("stopwords")

class LDAExtractor:
    def __init__(self, filepath, num_topics=10, words_per_topic=5):
        self.filepath = filepath
        self.num_topics = num_topics
        self.words_per_topic = words_per_topic
        self.paragraphs = []
        self.cleaned = []
        self.topics = []

    def load_docx(self):
        doc = docx.Document(self.filepath)
        self.paragraphs = [
            p.text.strip()
            for p in doc.paragraphs
            if len(p.text.strip()) > 30
        ]

    def preprocess(self):
        stop_words = set(stopwords.words("english"))
        for para in self.paragraphs:
            txt = para.lower()
            txt = re.sub(r"[^a-z\s]", " ", txt)
            tokens = [w for w in txt.split() if w not in stop_words]
            if len(tokens) > 3:
                self.cleaned.append(" ".join(tokens))

    def vectorize_and_fit(self):
        vec = CountVectorizer(max_df=0.85, min_df=2, stop_words="english")
        dtm = vec.fit_transform(self.cleaned)
        lda = LatentDirichletAllocation(n_components=self.num_topics,
                                        random_state=42,
                                        learning_method="batch")
        lda.fit(dtm)
        return lda, dtm, vec.get_feature_names_out()

    def extract_topics(self, lda, feature_names):
        for t in lda.components_:
            top_idxs = t.argsort()[-self.words_per_topic:]
            words = [feature_names[i] for i in top_idxs]
            title = " ".join(words[::-1]).title()
            self.topics.append(title)

    async def run_pipeline(self):
        self.load_docx()
        self.preprocess()
        lda, dtm, terms = self.vectorize_and_fit()
        self.extract_topics(lda, terms)
        return self.topics
